from django.shortcuts import render, redirect
from .models import RiskAssessment, Control, AssessmentControl
from django.http import HttpResponse
from django.template.loader import render_to_string
import pdfkit
from docx import Document

def home(request):
    return render(request, 'risk_assessment/home.html')

def general_info(request):
    if request.method == 'POST':
        title = request.POST.get('title')
        description = request.POST.get('description')
        prepared_by = request.POST.get('prepared_by')
        date = request.POST.get('date')
        
        request.session['general_info'] = {
            'title': title,
            'description': description,
            'prepared_by': prepared_by,
            'date': date
        }
        
        return redirect('asset_assessments')

    return render(request, 'risk_assessment/general_info.html')

def asset_assessments(request):
    general_info = request.session.get('general_info', {})
    asset_assessments = request.session.get('asset_assessments', [])

    if request.method == 'POST':
        assets = request.POST.getlist('assets[]')
        asset_categories = request.POST.getlist('asset_categories[]')
        criticalities = request.POST.getlist('criticalities[]')

        request.session['asset_assessments'] = [
            {'asset': asset, 'category': category, 'criticality': criticality}
            for asset, category, criticality in zip(assets, asset_categories, criticalities)
        ]
        
        return redirect('threat_hazard_assessment')
    
    return render(request, 'risk_assessment/asset_assessments.html', {
        'general_info': general_info,
        'asset_assessments': asset_assessments
    })

def save_asset_criteria(request):
    if request.method == 'POST':
        criticality_rating = request.POST.get('criticality_rating')
        asset_index = int(request.POST.get('asset_index'))
        
        asset_assessments = request.session.get('asset_assessments', [])
        if asset_index < len(asset_assessments):
            asset_assessments[asset_index]['criticality'] = criticality_rating
            request.session['asset_assessments'] = asset_assessments
        
        return redirect('asset_assessments')
    return redirect('general_info')

def threat_hazard_assessment(request):
    general_info = request.session.get('general_info', {})
    asset_assessments = request.session.get('asset_assessments', [])
    return render(request, 'risk_assessment/threat_hazard_assessment.html', {
        'general_info': general_info,
        'asset_assessments': asset_assessments
    })

def save_threat_hazard_assessment(request):
    if request.method == 'POST':
        threats_hazards = request.POST.getlist('threats_hazards[]')
        events = request.POST.getlist('events[]')
        assets = request.POST.getlist('assets[]')
        ratings = request.POST.getlist('ratings[]')

        threat_hazard_assessments = [
            {'threat': threat_hazard, 'event': event, 'asset': asset, 'rating': rating}
            for threat_hazard, event, asset, rating in zip(threats_hazards, events, assets, ratings)
        ]

        request.session['threat_hazard_assessments'] = threat_hazard_assessments
        
        return redirect('controls_assessment')
    return redirect('threat_hazard_assessment')

def controls_assessment(request):
    general_info = request.session.get('general_info', {})
    asset_assessments = request.session.get('asset_assessments', [])
    controls = Control.objects.all()
    return render(request, 'risk_assessment/controls_assessment.html', {
        'general_info': general_info,
        'asset_assessments': asset_assessments,
        'controls': controls
    })

def save_controls_assessment(request):
    if request.method == 'POST':
        controls = request.POST.getlist('controls[]')
        descriptions = request.POST.getlist('descriptions[]')
        effectiveness = request.POST.getlist('effectiveness[]')

        risk_assessment_id = request.session.get('risk_assessment_id')
        risk_assessment = RiskAssessment.objects.get(id=risk_assessment_id)

        for control_name, description, effectiveness_rating in zip(controls, descriptions, effectiveness):
            control, created = Control.objects.get_or_create(
                name=control_name,
                description=description,
                effectiveness=effectiveness_rating
            )
            AssessmentControl.objects.create(risk_assessment=risk_assessment, control=control)

        return redirect('threat_tool')
    return redirect('controls_assessment')

def save_changes(request):
    if request.method == 'POST':
        return redirect('home')
    else:
        return redirect('general_info')

def save_intent(request):
    if request.method == 'POST':
        none_intent = request.POST.get('none_intent')
        implied_intent = request.POST.get('implied_intent')
        expressed_intent = request.POST.get('expressed_intent')
        
        intent_data = {
            'none_intent': none_intent,
            'implied_intent': implied_intent,
            'expressed_intent': expressed_intent
        }
        
        request.session['intent_data'] = intent_data
        
        return redirect('general_info')
    return redirect('general_info')

def save_capability(request):
    if request.method == 'POST':
        extensive_capability = request.POST.get('extensive_capability')
        moderate_capability = request.POST.get('moderate_capability')
        low_capability = request.POST.get('low_capability')
        
        capability_data = {
            'extensive_capability': extensive_capability,
            'moderate_capability': moderate_capability,
            'low_capability': low_capability
        }
        
        request.session['capability_data'] = capability_data
        
        return redirect('general_info')
    return redirect('general_info')

def save_tolerance(request):
    if request.method == 'POST':
        threat_tolerance = request.POST.get('threat_tolerance')
        
        tolerance_data = {
            'threat_tolerance': threat_tolerance
        }
        
        request.session['tolerance_data'] = tolerance_data
        
        return redirect('general_info')
    return redirect('general_info')

def save_matrix(request):
    if request.method == 'POST':
        matrix_data = request.POST.getlist('matrix_data[]')
        
        request.session['matrix_data'] = matrix_data
        
        return redirect('general_info')
    return redirect('general_info')

def save_hazard_criteria(request):
    if request.method == 'POST':
        critical_criteria = request.POST.get('critical_criteria')
        emerging_criteria = request.POST.get('emerging_criteria')
        benign_criteria = request.POST.get('benign_criteria')
        event_tolerance = request.POST.get('event_tolerance')
        
        hazard_criteria = {
            'critical_criteria': critical_criteria,
            'emerging_criteria': emerging_criteria,
            'benign_criteria': benign_criteria,
            'event_tolerance': event_tolerance
        }
        
        request.session['hazard_criteria'] = hazard_criteria
        
        return redirect('general_info')
    return redirect('general_info')

def save_control_effectiveness(request):
    if request.method == 'POST':
        adequate_criteria = request.POST.get('adequate_criteria')
        moderate_criteria = request.POST.get('moderate_criteria')
        inadequate_criteria = request.POST.get('inadequate_criteria')
        
        control_effectiveness_data = {
            'adequate_criteria': adequate_criteria,
            'moderate_criteria': moderate_criteria,
            'inadequate_criteria': inadequate_criteria
        }
        
        request.session['control_effectiveness_data'] = control_effectiveness_data
        
        return redirect('general_info')
    return redirect('general_info')

def threat_tool(request):
    general_info = request.session.get('general_info', {})
    threat_hazard_assessments = request.session.get('threat_hazard_assessments', [])
    risk_matrix_options = [
        'Low',
        'Medium',
        'High',
        'Critical'
    ] # Replace with actual options from your Risk Matrix

    return render(request, 'risk_assessment/threat_tool.html', {
        'general_info': general_info,
        'threat_hazard_assessments': threat_hazard_assessments,
        'risk_matrix_options': risk_matrix_options
    })

def completed_assessment(request):
    general_info = request.session.get('general_info', {})
    asset_assessments = request.session.get('asset_assessments', [])
    threat_hazard_assessments = request.session.get('threat_hazard_assessments', [])
    controls = Control.objects.all()

    if request.method == 'POST':
        # Add logic to save the edited data if required
        pass

    return render(request, 'risk_assessment/completed_assessment.html', {
        'general_info': general_info,
        'asset_assessments': asset_assessments,
        'threat_hazard_assessments': threat_hazard_assessments,
        'controls': controls
    })

def download_pdf(request):
    general_info = request.session.get('general_info', {})
    asset_assessments = request.session.get('asset_assessments', [])
    threat_hazard_assessments = request.session.get('threat_hazard_assessments', [])
    controls = Control.objects.all()

    html = render_to_string('risk_assessment/completed_assessment_pdf.html', {
        'general_info': general_info,
        'asset_assessments': asset_assessments,
        'threat_hazard_assessments': threat_hazard_assessments,
        'controls': controls
    })
    pdf = pdfkit.from_string(html, False)
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="completed_assessment.pdf"'
    return response

def download_word(request):
    general_info = request.session.get('general_info', {})
    asset_assessments = request.session.get('asset_assessments', [])
    threat_hazard_assessments = request.session.get('threat_hazard_assessments', [])
    controls = Control.objects.all()

    document = Document()
    document.add_heading('Completed Assessment', 0)

    document.add_heading('General Info', level=1)
    document.add_paragraph(f"Title: {general_info.get('title')}")
    document.add_paragraph(f"Description: {general_info.get('description')}")
    document.add_paragraph(f"Prepared By: {general_info.get('prepared_by')}")
    document.add_paragraph(f"Date: {general_info.get('date')}")

    document.add_heading('Asset Assessments', level=1)
    for assessment in asset_assessments:
        document.add_paragraph(f"Asset: {assessment['asset']}, Category: {assessment['category']}, Criticality: {assessment['criticality']}")

    document.add_heading('Threat/Hazard Assessments', level=1)
    for assessment in threat_hazard_assessments:
        document.add_paragraph(f"Threat: {assessment['threat']}, Event: {assessment['event']}, Asset: {assessment['asset']}, Rating: {assessment['rating']}")

    document.add_heading('Controls', level=1)
    for control in controls:
        document.add_paragraph(f"Control: {control.name}, Description: {control.description}, Effectiveness: {control.effectiveness}")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="completed_assessment.docx"'
    document.save(response)
    return response

def custom_404(request, exception):
    return render(request, 'risk_assessment/404.html', status=404)

def custom_500(request):
    return render(request, 'risk_assessment/500.html', status=500)
